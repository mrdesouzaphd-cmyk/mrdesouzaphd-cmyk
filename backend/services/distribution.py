"""Multi-platform ebook distribution service.

Handles publishing to:
- Amazon Kindle Direct Publishing (KDP)
- Audiobook platforms (ACX/Audible)
- Gumroad
- Other platforms via generic API integration

Also handles export formats compatible with:
- Apple Pages (DOCX export)
- Canva (structured content JSON)
- Adobe Express (asset preparation)
"""

import json
import structlog
from dataclasses import dataclass

from google.cloud import texttospeech

from backend.services.storage import StorageService
from config.settings import settings

logger = structlog.get_logger()


@dataclass
class DistributionResult:
    platform: str
    status: str  # "success", "pending", "error"
    url: str | None = None
    message: str = ""


class DistributionService:
    """Manages ebook distribution to multiple sales platforms."""

    def __init__(self, storage_service: StorageService):
        self.storage = storage_service

    # ──────────────────────────────────────────
    # Amazon Kindle Direct Publishing
    # ──────────────────────────────────────────

    def prepare_for_kindle(self, project, ebook_gcs_path: str) -> dict:
        """Prepare ebook metadata and file for Amazon KDP upload.

        KDP requires:
        - EPUB or DOCX file
        - Book metadata (title, description, keywords, categories)
        - Cover image (2560x1600 px recommended)

        Returns a package dict with all required assets and metadata.
        """
        metadata = {
            "title": project.title,
            "description": project.description or "",
            "language": "en" if project.language == "en" else "pt",
            "keywords": [],  # To be populated from AI analysis
            "categories": [],
            "author": "DrDeSouzAI",
            "publisher": "",
            "isbn": "",  # User provides
        }

        # Download the ebook file for local packaging
        ebook_data = self.storage.download_file(ebook_gcs_path)

        # Generate KDP-ready package path
        package_path = self.storage.upload_processed_content(
            project_id=str(project.id),
            file_data=ebook_data,
            filename=f"kdp_{project.title.replace(' ', '_')}.epub",
            content_type="application/epub+zip",
        )

        logger.info("kdp_package_prepared", project_id=str(project.id))

        return {
            "platform": "amazon_kdp",
            "metadata": metadata,
            "ebook_gcs_path": package_path,
            "instructions": (
                "Upload the EPUB file to kdp.amazon.com along with the metadata. "
                "KDP does not currently offer a public API for automated uploads — "
                "use the KDP web dashboard to complete publishing."
            ),
        }

    # ──────────────────────────────────────────
    # Audiobook Generation
    # ──────────────────────────────────────────

    def generate_audiobook(
        self,
        project,
        text_content: str,
        language: str = "en",
        voice_name: str | None = None,
    ) -> str:
        """Generate an audiobook from text using Google Cloud Text-to-Speech.

        Returns GCS path to the generated audio file.
        """
        client = texttospeech.TextToSpeechClient()

        # Select voice based on language
        if voice_name is None:
            voice_name = "en-US-Studio-O" if language == "en" else "pt-BR-Standard-B"

        voice = texttospeech.VoiceSelectionParams(
            language_code="en-US" if language == "en" else "pt-BR",
            name=voice_name,
        )

        audio_config = texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3,
            speaking_rate=0.95,  # Slightly slower for audiobook clarity
            pitch=0.0,
        )

        # Split text into chunks (TTS has character limits)
        chunks = self._split_text_for_tts(text_content, max_chars=4500)
        audio_parts = []

        for i, chunk in enumerate(chunks):
            synthesis_input = texttospeech.SynthesisInput(text=chunk)
            response = client.synthesize_speech(
                input=synthesis_input,
                voice=voice,
                audio_config=audio_config,
            )
            audio_parts.append(response.audio_content)
            logger.info("tts_chunk_generated", chunk=i + 1, total=len(chunks))

        # Concatenate audio
        full_audio = b"".join(audio_parts)

        # Upload to GCS
        gcs_path = self.storage.upload_ebook(
            project_id=str(project.id),
            file_data=full_audio,
            filename=f"{project.title.replace(' ', '_')}_audiobook.mp3",
            content_type="audio/mpeg",
        )

        logger.info(
            "audiobook_generated",
            project_id=str(project.id),
            size_mb=len(full_audio) / (1024 * 1024),
        )

        return gcs_path

    def prepare_audiobook_script(self, content_items) -> str:
        """Extract and format text from content items as an audiobook narration script.

        Handles:
        - Text chapters: direct inclusion
        - Images: reads the AI caption/description as narration
        - Videos: reads the AI summary as narration
        """
        script_parts = []

        for item in content_items:
            title = item.title or "Untitled Section"
            script_parts.append(f"\n\n{title}.\n\n")

            if item.content_type == "text":
                script_parts.append(item.raw_text or "")
            elif item.content_type == "image":
                if item.ai_caption:
                    script_parts.append(item.ai_caption)
                if item.ai_description:
                    script_parts.append(f" {item.ai_description}")
            elif item.content_type == "video":
                if item.ai_description:
                    script_parts.append(item.ai_description)

        return "\n".join(script_parts)

    # ──────────────────────────────────────────
    # Export for Apple Pages / Canva / Adobe
    # ──────────────────────────────────────────

    def export_for_pages(self, project, content_items) -> str:
        """Export content as a DOCX file compatible with Apple Pages.

        Pages can open DOCX files and apply its templates to them.
        This gives you the best of both worlds: automated content +
        Pages template design.
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        doc = Document()

        # Title page
        title_para = doc.add_heading(project.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if project.description:
            desc_para = doc.add_paragraph(project.description)
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Chapters
        for item in content_items:
            chapter_title = item.title or "Untitled"
            doc.add_heading(chapter_title, level=1)

            if item.content_type == "text" and item.raw_text:
                for paragraph in item.raw_text.split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())

            elif item.content_type == "image":
                if item.gcs_processed_path or item.gcs_raw_path:
                    try:
                        path = item.gcs_processed_path or item.gcs_raw_path
                        img_data = self.storage.download_file(path)
                        img_stream = BytesIO(img_data)
                        doc.add_picture(img_stream, width=Inches(4.5))
                    except Exception:
                        doc.add_paragraph(f"[Image: {chapter_title}]")
                if item.ai_caption:
                    caption = doc.add_paragraph(item.ai_caption)
                    caption.runs[0].italic = True if caption.runs else None

            elif item.content_type == "video":
                doc.add_paragraph(
                    f"[Video Content: {chapter_title}]",
                    style="Intense Quote",
                )
                if item.ai_description:
                    doc.add_paragraph(item.ai_description)

            doc.add_page_break()

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        # Upload to GCS
        gcs_path = self.storage.upload_ebook(
            project_id=str(project.id),
            file_data=docx_bytes,
            filename=f"{project.title.replace(' ', '_')}_pages.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        logger.info("pages_export_created", project_id=str(project.id))
        return gcs_path

    def export_for_canva(self, project, content_items) -> str:
        """Export content as structured JSON for Canva template population.

        Canva's API allows populating templates with structured data.
        This exports content in a format ready for Canva integration.
        """
        canva_data = {
            "project_title": project.title,
            "project_description": project.description or "",
            "language": project.language,
            "pages": [],
        }

        # Cover page
        canva_data["pages"].append({
            "type": "cover",
            "title": project.title,
            "subtitle": project.description or "",
        })

        for item in content_items:
            page = {
                "type": item.content_type,
                "title": item.title or "Untitled",
                "order": item.chapter_order,
            }

            if item.content_type == "text":
                page["body"] = item.raw_text or ""
            elif item.content_type == "image":
                page["image_path"] = item.gcs_processed_path or item.gcs_raw_path or ""
                page["caption"] = item.ai_caption or ""
                page["alt_text"] = item.ai_description or ""
            elif item.content_type == "video":
                page["summary"] = item.ai_description or ""
                page["caption"] = item.ai_caption or ""

            canva_data["pages"].append(page)

        json_bytes = json.dumps(canva_data, indent=2, ensure_ascii=False).encode("utf-8")

        gcs_path = self.storage.upload_ebook(
            project_id=str(project.id),
            file_data=json_bytes,
            filename=f"{project.title.replace(' ', '_')}_canva.json",
            content_type="application/json",
        )

        logger.info("canva_export_created", project_id=str(project.id))
        return gcs_path

    # ──────────────────────────────────────────
    # Gumroad Distribution
    # ──────────────────────────────────────────

    def prepare_for_gumroad(self, project, ebook_gcs_path: str, price: float) -> dict:
        """Prepare ebook for Gumroad listing.

        Gumroad has an API for creating products programmatically.
        """
        return {
            "platform": "gumroad",
            "product": {
                "name": project.title,
                "description": project.description or "",
                "price": int(price * 100),  # Gumroad uses cents
                "currency": "usd",
                "file_gcs_path": ebook_gcs_path,
            },
            "instructions": (
                "Use the Gumroad API (POST /products) with your access token "
                "to create the product listing. The file will need to be uploaded "
                "separately via their file upload endpoint."
            ),
        }

    # ──────────────────────────────────────────
    # Helpers
    # ──────────────────────────────────────────

    def _split_text_for_tts(self, text: str, max_chars: int = 4500) -> list[str]:
        """Split text into chunks suitable for TTS API limits.

        Splits at paragraph boundaries to maintain natural pauses.
        """
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) + 2 > max_chars:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para
            else:
                current_chunk += "\n\n" + para

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks if chunks else [text[:max_chars]]
